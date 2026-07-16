from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('VectorDBGen 개발 문서 (시스템 구조 및 프로세스)', 0)

doc.add_heading('1. 시스템 개요', level=1)
doc.add_paragraph('VectorDBGen은 AI-AutoRouting 프로젝트에서 사용되는 C# WPF 애플리케이션으로, 벡터 데이터베이스(PostgreSQL + pgvector)의 스키마를 초기화하고 파이썬 기반의 데이터 파이프라인 엔진을 구동하여 벡터를 생성/적재하는 도구입니다. 데이터베이스 상태 모니터링, 외부 파이썬 프로세스와의 비동기 연동, 엑셀 형태의 DDL 실행 결과 리포팅 기능을 지원합니다.')

doc.add_heading('2. 전체 프로세스 흐름', level=1)
p = doc.add_paragraph()
p.add_run('1. DB 접속 및 상태 모니터링: ').bold = True
p.add_run('사용자가 호스트, 포트, DB명 등을 입력하고 연결을 요청하면, Npgsql을 통해 DB 커넥션을 맺고 연결 상태를 갱신합니다.\n')
p.add_run('2. 스키마 초기화 (DDL 실행): ').bold = True
p.add_run('Feature Vector, Context Vector, Auto Design(Group/Segment) 등 각 벡터 테이블 생성을 위한 SQL 스크립트(.sql)를 불러와 DB에 실행하고, 테이블을 생성합니다.\n')
p.add_run('3. 의존성 사전 검증 (Preflight Check): ').bold = True
p.add_run('벡터를 추출하기 위해 선행되어야 하는 원본 테이블들(예: TB_ROUTE_PATH, TB_BIM_OBSTACLE 등)이 DB에 존재하는지 information_schema를 조회하여 유효성을 검증합니다.\n')
p.add_run('4. 벡터 생성 엔진 구동 (Python Interop): ').bold = True
p.add_run('검증을 통과하면, C# Process API를 이용해 지정된 파이썬 스크립트(BuildFeatureVectors.py 등)를 실행하며, 생성되는 실시간 로그(stdout)를 UI의 텍스트 박스로 스트리밍합니다.\n')
p.add_run('5. 결과 보고 및 DDL 리포트 추출: ').bold = True
p.add_run('전체 과정이 종료되면 ClosedXML을 이용해 성공/실패 여부를 엑셀 형태(.xlsx)로 내보냅니다.')

doc.add_heading('3. 핵심 변수 및 데이터 구조', level=1)
p2 = doc.add_paragraph()
p2.add_run('_builderSourceRequirements: ').bold = True
p2.add_run('각 파이프라인 단계별로 요구되는 원본 테이블명들을 매핑해놓은 Dictionary 구조입니다. Preflight Check 단계에서 이 목록을 기준으로 검증을 수행합니다.\n')
p2.add_run('_builderTargetTable: ').bold = True
p2.add_run('파이썬 엔진이 데이터를 적재할 타겟 테이블명을 정의한 변수로, 실행 전 TRUNCATE 또는 DELETE 작업 시 활용됩니다.\n')
p2.add_run('_builderDdlFile: ').bold = True
p2.add_run('스키마 초기화를 위한 SQL 스크립트 파일명을 관리합니다.\n')
p2.add_run('_proc: ').bold = True
p2.add_run('System.Diagnostics.Process 객체로, 파이썬 파이프라인을 비동기로 띄우고 라이프사이클을 관리합니다.')

doc.add_heading('4. 주요 함수 상세', level=1)
f1 = doc.add_paragraph(style='List Bullet')
f1.add_run('CheckConnectionAsync(): ').bold = True
f1.add_run('DB 접속 가능 여부를 확인하고, 실패 시 UI에 에러를 표기합니다.')
f2 = doc.add_paragraph(style='List Bullet')
f2.add_run('EnsureFreshTargetTableAsync(string tag): ').bold = True
f2.add_run('벡터 데이터 적재 전 타겟 테이블을 비우는 역할을 수행합니다. TRUNCATE RESTART IDENTITY를 시도하며, FK 제약조건으로 인해 실패 시 DELETE FROM을 시도합니다.')
f3 = doc.add_paragraph(style='List Bullet')
f3.add_run('PreflightCheckAsync(string tag): ').bold = True
f3.add_run('해당 태그(feature, context 등)에 정의된 원본 테이블이 information_schema 상에 존재하는지 확인하여, 누락된 경우 작업을 차단합니다.')
f4 = doc.add_paragraph(style='List Bullet')
f4.add_run('RunPythonAsync(string cwd, string script, string args, CancellationToken ct): ').bold = True
f4.add_run('PYTHONUNBUFFERED=1 환경변수를 설정하여 파이썬 프로세스를 비동기 실행합니다. 표준 출력과 표준 에러를 비동기 이벤트(BeginOutputReadLine)로 캡처하여 UI 스레드로 릴레이(Dispatcher.Invoke)합니다.')
f5 = doc.add_paragraph(style='List Bullet')
f5.add_run('RunDdlAsync(string ddlFile): ').bold = True
f5.add_run('SQL 파일을 읽어들여 명령문(;) 단위로 분할하여 순차적으로 NpgsqlCommand로 실행하고, 각 쿼리의 성공/실패 건수를 집계합니다.')
f6 = doc.add_paragraph(style='List Bullet')
f6.add_run('BtnRunAll_Click(): ').bold = True
f6.add_run('Feature -> Context -> Group -> Segment 의 4단계 파이프라인을 순차적으로 호출하며, 한 단계라도 에러 발생 시 이후 단계를 스킵(Skip)하는 오케스트레이션 역할을 담당합니다.')

doc.add_heading('5. 주요 알고리즘 및 구현 특징', level=1)
a1 = doc.add_paragraph()
a1.add_run('실시간 스트리밍 아키텍처: ').bold = True
a1.add_run('파이썬 엔진의 연산 시간이 길어짐에 따라 GUI 스레드가 Block되는 현상을 막기 위해 async/await 패턴과 CancellationTokenSource를 도입했습니다. UI 업데이트는 DispatcherTimerLite와 Process.OutputDataReceived 이벤트를 통해 매우 가벼운 부하로 로그창에 출력되도록 구현되었습니다.\n\n')
a1.add_run('Fallback 스크립트 탐색 체계 (ResolveScriptDir): ').bold = True
a1.add_run('배포 환경 혹은 개발 환경의 위치가 달라질 수 있으므로, 현재 실행 디렉토리(AppDomain.CurrentDomain.BaseDirectory)에서 시작하여 상위 경로를 탐색하고, 최종적으로 원본 소스 디렉토리(RoutingAI/src)로 fallback 하는 로직이 적용되어 환경 독립성을 높였습니다.')

doc.save('D:/DINNO/DEV/AI-AutoRouting/TopKGen/Docs/VectorDBGen_Architecture.docx')
