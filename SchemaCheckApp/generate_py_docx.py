from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
doc.add_heading('VectorDBGen Python 파이프라인 엔진 명세서', 0)

doc.add_heading('1. 개요', level=1)
doc.add_paragraph('본 문서는 VectorDBGen에서 내부적으로 호출하는 4단계 파이썬 벡터 데이터 생성 파이프라인의 핵심 기능, 알고리즘, 함수 및 변수 구성을 상세히 설명합니다.')

# --- BuildFeatureVectors.py ---
doc.add_heading('2. BuildFeatureVectors.py (1단계: Feature Vector 30D)', level=1)
doc.add_paragraph('기하학적 경로 데이터를 로드하여 30차원 Feature Vector를 생성하고 데이터베이스(TB_ROUTE_FEATURE_VECTOR)에 적재합니다.')

doc.add_heading('주요 알고리즘', level=2)
doc.add_paragraph('경로의 기하학적 형태를 30차원의 밀집 벡터(Dense Vector)로 임베딩합니다. 시작/종료 위상(방향 벡터), 바운딩 박스 크기 비율, 그리고 구간별(33%씩) 꺾임 횟수 등 물리적 특성을 정규화(Min-Max Normalization)하여 반영합니다.')

doc.add_heading('주요 함수', level=2)
f1 = doc.add_paragraph(style='List Bullet')
f1.add_run('main(): ').bold = True
f1.add_run('CLI 인자(from-db, local 등)를 파싱하고 DB 커넥션을 맺어 파이프라인을 트리거합니다.')
f2 = doc.add_paragraph(style='List Bullet')
f2.add_run('build_vectors(records): ').bold = True
f2.add_run('경로 데이터 배열을 순회하며 TopKRoutingSearch.RoutePathEncoder 클래스를 호출하여 각 경로의 30D 벡터 및 방향 패턴(Direction Pattern)을 추출합니다.')
f3 = doc.add_paragraph(style='List Bullet')
f3.add_run('save_vectors_db(db_params, vectors): ').bold = True
f3.add_run('psycopg2의 execute_values를 활용해 생성된 30D 벡터 배열과 메타데이터를 TB_ROUTE_FEATURE_VECTOR 테이블에 고속 삽입(Bulk Insert)합니다.')

# --- BuildContextVectors.py ---
doc.add_heading('3. BuildContextVectors.py (2단계: Context Vector 30D)', level=1)
doc.add_paragraph('TB_BIM_OBSTACLE을 BAY 범위로 조회하여 시작·종점의 500/1,000mm 장애물 환경 벡터(30D)를 추출합니다.')

doc.add_heading('주요 알고리즘', level=2)
doc.add_paragraph('시작·종점에서 AABB 표면거리 기준 0~500mm와 500~1,000mm shell을 탐색합니다. 기둥·보 개수, 최근접 표면거리·방향, free-space 및 Tier3 특징을 30차원 Context Vector로 병합합니다.')

doc.add_heading('주요 함수', level=2)
c1 = doc.add_paragraph(style='List Bullet')
c1.add_run('_load_route_points(cur, guid): ').bold = True
c1.add_run('경로 GUID에 속하는 모든 세그먼트의 시작/종료 3D 좌표를 로드합니다.')
c2 = doc.add_paragraph(style='List Bullet')
c2.add_run('main(): ').bold = True
c2.add_run('Feature Vector가 생성된 경로들을 대상으로 주변 환경 데이터를 병합 인코딩하고, TB_ROUTE_CONTEXT_VECTOR 테이블에 삽입합니다.')

# --- BuildDesignGroups.py ---
doc.add_heading('4. BuildDesignGroups.py (3단계: Design Group)', level=1)
doc.add_paragraph('유사한 기하학적 형태 및 맥락을 가진 경로들을 클러스터링하여 논리적인 그룹(Design Group)으로 묶습니다.')

doc.add_heading('주요 알고리즘', level=2)
doc.add_paragraph('생성된 30D Feature Vector들을 대상으로 K-Means 혹은 DBSCAN(유사도 임계값 기반) 클러스터링을 수행하여, 시작점/종료점/패턴이 유사한 배관들을 하나의 묶음(Group)으로 분류합니다. 분류된 그룹 정보는 TB_ROUTE_DESIGN_GROUP 테이블에 기록됩니다.')

doc.add_heading('주요 함수', level=2)
g1 = doc.add_paragraph(style='List Bullet')
g1.add_run('main(): ').bold = True
g1.add_run('AutoRouteDesigner의 group_builder 모듈을 호출하여 전체 경로를 대상으로 군집화를 수행하고 결과를 DB에 반영합니다.')

# --- BuildSegmentTemplates.py ---
doc.add_heading('5. BuildSegmentTemplates.py (4단계: Segment Templates)', level=1)
doc.add_paragraph('세부 구간(Segment) 레벨에서의 반복적인 배관 템플릿 패턴을 추출합니다.')

doc.add_heading('주요 알고리즘', level=2)
doc.add_paragraph('경로 내부의 단위 세그먼트들의 길이, 꺾임, 부속(Fitting) 조합을 분석하여 빈번하게 등장하는 마이크로 라우팅 패턴을 도출하고 TB_ROUTE_SEGMENT_TEMPLATE 테이블에 적재합니다. 이는 이후 국소 라우팅 단계에서 참조 템플릿으로 활용됩니다.')

doc.add_heading('주요 변수/상수 (전역 통용)', level=2)
v1 = doc.add_paragraph(style='List Bullet')
v1.add_run('db_params (dict): ').bold = True
v1.add_run('DB 호스트, 포트, DB명, 사용자 계정 등 psycopg2 연결에 필요한 정보를 담는 딕셔너리입니다. (DDW_AI_DB 접속용)')
v2 = doc.add_paragraph(style='List Bullet')
v2.add_run('norm_params (dict): ').bold = True
v2.add_run('Min-Max 정규화 과정에 사용되는 각 차원별 최소/최대값 파라미터가 담겨있으며, json 파일로 캐싱되어 추후 추론(Inference) 시 로드됩니다.')

doc.save('D:/DINNO/DEV/AI-AutoRouting/TopKGen/Docs/Python_Pipelines_Architecture.docx')
